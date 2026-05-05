"""
generate_report_docx.py - Generates an editable Word document (.docx) report
with embedded matplotlib graphs for the Alzheimer's Detection project.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
GRAPHS_DIR = os.path.join(SCRIPT_DIR, 'report_graphs')

# ── Styling helpers ──

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_graph_image(doc, filename, width=6.0):
    """Add a graph image from report_graphs folder."""
    path = os.path.join(GRAPHS_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # spacing


def build_docx():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading("Alzheimer's Disease Detection", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Using Deep Learning", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tagline = doc.add_paragraph("Comprehensive Project Report")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].font.size = Pt(16)
    tagline.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        "Domain: Healthcare / Medical Image Classification",
        "Dataset: OASIS Alzheimer's MRI (86,437 images)",
        "Best Model: MobileNetV2 (Fine-Tuned) - 99.47% Accuracy",
        "Platform: Google Colab (T4 GPU) + Flask Web Application",
        "Date: April 2026",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Project Overview",
        "2. Problem Statement & Purpose",
        "3. Dataset Description",
        "4. Data Splitting & Augmentation",
        "5. Model Architectures",
        "6. Experiment 1 - Initial Model Comparison",
        "7. Experiment 2 - Fine-Tuned MobileNetV2",
        "8. Performance Metrics",
        "9. Final Comparison & Improvement",
        "10. Technologies Used",
        "11. Web Application",
        "12. Conclusion",
        "13. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "1. Project Overview", level=1)
    add_body(doc,
        "This project is an end-to-end Alzheimer's disease detection system that uses deep learning models "
        "trained on brain MRI scans to classify the severity of Alzheimer's disease into four stages. "
        "The system includes a full machine learning pipeline for model training on Google Colab (T4 GPU) "
        "and a professional Flask web application for real-time prediction with clinical context."
    )
    doc.add_paragraph()
    add_styled_table(doc,
        ['Field', 'Details'],
        [
            ['Project Title', "Alzheimer's Disease Detection System"],
            ['Domain', 'Healthcare / Medical Image Classification'],
            ['Technique', 'Deep Learning (CNNs & Transfer Learning)'],
            ['Language', 'Python 3.12'],
            ['Framework', 'TensorFlow 2.19 / Keras'],
            ['Web App', 'Flask + SQLite + Bootstrap 5'],
            ['Training', 'Google Colab with NVIDIA T4 GPU'],
            ['Best Accuracy', '99.47% (MobileNetV2 Fine-Tuned)'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 2. PROBLEM STATEMENT
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "2. Problem Statement & Purpose", level=1)

    add_heading_styled(doc, "What is the Project?", level=2)
    add_body(doc,
        "This project develops an AI-powered system that can analyze brain MRI (Magnetic Resonance Imaging) "
        "scans and classify them into one of four stages of Alzheimer's disease progression: Non Demented, "
        "Very Mild Demented, Mild Demented, and Moderate Demented."
    )

    add_heading_styled(doc, "Why am I Doing It?", level=2)
    add_bullet(doc, "Alzheimer's disease affects over 55 million people worldwide. Early and accurate detection is critical.")
    add_bullet(doc, "Manual diagnosis by radiologists is time-intensive, subjective, and prone to inter-observer variability.")
    add_bullet(doc, "AI-assisted detection can serve as a second opinion tool, improving diagnostic accuracy.")
    add_bullet(doc, "The distinction between early stages is extremely subtle on MRI, making it ideal for deep learning.")

    add_heading_styled(doc, "What is the Purpose?", level=2)
    add_bullet(doc, "Build a robust multi-class image classifier across four Alzheimer's stages.")
    add_bullet(doc, "Compare 6 different deep learning architectures (CNN, MLP, VGG16, ResNet50, EfficientNetB0, MobileNetV2).")
    add_bullet(doc, "Deploy the best model in a user-friendly web application with clinical context.")
    add_bullet(doc, "Demonstrate that transfer learning with fine-tuning achieves near-perfect accuracy (99.47%).")

    add_heading_styled(doc, "How Did I Do It?", level=2)
    steps = [
        "Collected the OASIS brain MRI dataset (86,437 images across 4 classes).",
        "Cleaned the data - validated all images, removed corrupted/duplicate files.",
        "Split the data into Train (70%), Validation (15%), Test (15%) with stratification.",
        "Applied data augmentation and computed class weights for imbalance handling.",
        "Trained 6 different models on Google Colab with a T4 GPU.",
        "Identified MobileNetV2 as the best base architecture.",
        "Fine-tuned MobileNetV2 with 224x224 input and partial backbone unfreezing - achieved 99.47% accuracy.",
        "Built a Flask web application with user auth, upload, prediction, and clinical results.",
    ]
    for s in steps:
        add_bullet(doc, s)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 3. DATASET DESCRIPTION
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "3. Dataset Description", level=1)
    add_body(doc,
        "The dataset used is the OASIS (Open Access Series of Imaging Studies) Alzheimer's Detection Dataset, "
        "sourced from Kaggle (ninadaithal/imagesoasis). It contains 86,437 labeled brain MRI scan images "
        "distributed across 4 classes. The dataset exhibits severe class imbalance, with 'Non Demented' "
        "having 138x more samples than 'Moderate Dementia'."
    )
    doc.add_paragraph()
    add_styled_table(doc,
        ['Class', 'Images', 'Share (%)', 'Train', 'Test'],
        [
            ['Mild Dementia', '5,002', '5.79%', '3,501', '751'],
            ['Moderate Dementia', '488', '0.56%', '342', '73'],
            ['Non Demented', '67,222', '77.77%', '47,055', '10,083'],
            ['Very Mild Dementia', '13,725', '15.88%', '9,607', '2,059'],
            ['TOTAL', '86,437', '100%', '60,505', '12,966'],
        ]
    )
    doc.add_paragraph()
    add_graph_image(doc, '01_class_distribution.png', 6.2)
    add_graph_image(doc, '02_class_weights.png', 5.0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 4. DATA SPLITTING & AUGMENTATION
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "4. Data Splitting & Augmentation", level=1)
    add_body(doc,
        "The dataset was split using stratified sampling to maintain class proportions: "
        "70% for training (60,505 images), 15% for validation (12,966), and 15% for testing (12,966). "
        "Real-time data augmentation was applied during training to improve generalization."
    )
    add_graph_image(doc, '03_data_split.png', 5.5)

    add_heading_styled(doc, "Augmentation Parameters", level=2)
    add_styled_table(doc,
        ['Parameter', 'Experiment 1', 'Experiment 2 (Fine-Tuned)'],
        [
            ['Rotation', '20 degrees', '10 degrees'],
            ['Zoom', '20%', '10%'],
            ['Horizontal Flip', 'Yes', 'No (MRI-specific)'],
            ['Shear', '20%', '10%'],
            ['Normalization', '1/255', '1/255'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 5. MODEL ARCHITECTURES
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "5. Model Architectures", level=1)
    add_body(doc,
        "Six deep learning architectures were evaluated. Two are trained from scratch (CNN, MLP) and "
        "four use transfer learning from ImageNet pre-trained weights (VGG16, ResNet50, EfficientNetB0, MobileNetV2)."
    )

    models = [
        ("1. Custom CNN",
         "3 Conv2D blocks (32 > 64 > 128 filters) with MaxPooling, BatchNormalization, and Dropout, "
         "followed by Dense(256) and Softmax(4). Approximately 8.5M parameters."),
        ("2. MLP (Multi-Layer Perceptron)",
         "Fully connected network: Flatten > Dense(512) > Dense(256) > Dense(128) > Softmax(4). "
         "No convolutions. Approximately 25.3M parameters."),
        ("3. VGG16 (Transfer Learning)",
         "ImageNet pre-trained VGG16 backbone (all layers frozen) + GlobalAveragePooling2D + "
         "Dense(256) + Dropout(0.5) + Softmax(4). Approximately 14.8M total parameters."),
        ("4. ResNet50 (Transfer Learning)",
         "ImageNet pre-trained ResNet50 backbone (all layers frozen) + GlobalAveragePooling2D + "
         "Dense(256) + Dropout(0.5) + Softmax(4). Approximately 23.7M total parameters."),
        ("5. EfficientNetB0 (Transfer Learning)",
         "ImageNet pre-trained EfficientNetB0 backbone (all layers frozen) + GlobalAveragePooling2D + "
         "Dense(256) + Dropout(0.5) + Softmax(4). Approximately 4.1M total parameters."),
        ("6. MobileNetV2 (Fine-Tuned) [BEST MODEL]",
         "ImageNet pre-trained MobileNetV2 with last 30 layers unfrozen for fine-tuning + "
         "GlobalAveragePooling2D + Dense(256, ReLU) + Dropout(0.5) + Softmax(4). "
         "Total: 2,586,948 parameters. Trainable: 1,855,364 parameters (7.08 MB)."),
    ]
    for title, desc in models:
        add_heading_styled(doc, title, level=2)
        add_body(doc, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 6. EXPERIMENT 1
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "6. Experiment 1 - Initial Model Comparison", level=1)
    add_body(doc,
        "All six models were trained on the same dataset split with 128x128 input resolution, fully frozen "
        "pre-trained backbones, Adam optimizer (lr=1e-4), and 20 epochs. Only transfer learning models with "
        "sufficient feature extraction capability achieved meaningful accuracy."
    )
    doc.add_paragraph()
    add_styled_table(doc,
        ['Model', 'Test Accuracy', 'Weighted F1', 'Macro F1', 'Status'],
        [
            ['MobileNetV2', '75.27%', '0.78', '0.63', 'Best'],
            ['VGG16', '69.10%', '0.73', '0.44', 'Moderate'],
            ['ResNet50', '47.29%', '0.56', '0.26', 'Poor'],
            ['Custom CNN', '0.56%', '0.00', '0.00', 'Failed'],
            ['MLP', '0.56%', '0.00', '0.00', 'Failed'],
            ['EfficientNetB0', '0.56%', '0.00', '0.00', 'Failed'],
        ]
    )
    doc.add_paragraph()
    add_graph_image(doc, '04_exp1_comparison.png', 6.0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 7. EXPERIMENT 2
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "7. Experiment 2 - Fine-Tuned MobileNetV2 (224x224)", level=1)
    add_body(doc,
        "Based on Experiment 1 results, MobileNetV2 was selected for fine-tuning with higher resolution "
        "input (224x224), partial backbone unfreezing (last 30 layers trainable), lower learning rate (1e-4), "
        "and MRI-specific augmentation (no horizontal flip). This achieved 99.47% test accuracy."
    )

    add_heading_styled(doc, "Training Progression", level=2)
    add_styled_table(doc,
        ['Epoch', 'Train Acc', 'Val Acc', 'Train Loss', 'Val Loss'],
        [
            ['1', '70.91%', '83.87%', '0.6249', '0.3756'],
            ['2', '85.49%', '82.22%', '0.2453', '0.5224'],
            ['3', '90.73%', '92.91%', '0.1426', '0.1958'],
            ['4', '92.76%', '96.67%', '0.1196', '0.0895'],
            ['5', '93.93%', '96.83%', '0.1102', '0.0884'],
        ]
    )
    doc.add_paragraph()
    add_graph_image(doc, '05_mobilenet_training.png', 6.2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 8. PERFORMANCE METRICS
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "8. Performance Metrics", level=1)

    add_heading_styled(doc, "Classification Report - MobileNetV2 (Fine-Tuned)", level=2)
    add_body(doc, "Overall Test Accuracy: 99.47% on 12,966 test images.")
    doc.add_paragraph()
    add_styled_table(doc,
        ['Class', 'Precision', 'Recall', 'F1-Score', 'Support'],
        [
            ['Mild Dementia', '0.98', '1.00', '0.99', '751'],
            ['Moderate Dementia', '1.00', '1.00', '1.00', '73'],
            ['Non Demented', '1.00', '1.00', '1.00', '10,083'],
            ['Very Mild Dementia', '0.99', '0.98', '0.99', '2,059'],
            ['Weighted Average', '0.99', '0.99', '0.99', '12,966'],
        ]
    )
    doc.add_paragraph()
    add_graph_image(doc, '07_classification_report.png', 6.0)

    add_heading_styled(doc, "Confusion Matrix", level=2)
    add_graph_image(doc, '06_confusion_matrix.png', 5.0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 9. FINAL COMPARISON
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "9. Final Comparison & Improvement", level=1)
    add_graph_image(doc, '08_final_comparison.png', 6.2)
    add_graph_image(doc, '09_improvement.png', 5.8)

    add_heading_styled(doc, "Model Efficiency: Parameters vs Accuracy", level=2)
    add_body(doc,
        "MobileNetV2 achieved the highest accuracy (99.47%) with the fewest parameters (2.6M), "
        "demonstrating its superior efficiency for medical imaging tasks."
    )
    add_graph_image(doc, '10_params_vs_accuracy.png', 5.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 10. TECHNOLOGIES USED
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "10. Technologies Used", level=1)
    add_styled_table(doc,
        ['Category', 'Technology', 'Purpose'],
        [
            ['Language', 'Python 3.12', 'Core programming language'],
            ['Deep Learning', 'TensorFlow 2.19 / Keras', 'Model building and training'],
            ['Pre-trained', 'MobileNetV2 (ImageNet)', 'Transfer learning backbone'],
            ['Image Processing', 'OpenCV, PIL', 'Loading and resizing MRI scans'],
            ['Data Science', 'NumPy, Pandas', 'Array operations, dataframes'],
            ['Visualization', 'Matplotlib, Seaborn', 'Plots, confusion matrices'],
            ['ML Utilities', 'scikit-learn', 'Splitting, metrics, class weights'],
            ['Web Framework', 'Flask', 'Backend web application'],
            ['Database', 'SQLite3', 'User authentication, prediction logs'],
            ['Security', 'Werkzeug', 'Password hashing (PBKDF2-SHA256)'],
            ['Frontend', 'HTML5, CSS3, JS', 'UI templates and interactivity'],
            ['UI Framework', 'Bootstrap 5', 'Responsive design'],
            ['Training GPU', 'Google Colab (T4)', 'GPU-accelerated training'],
            ['Dataset Source', 'Kaggle (kagglehub)', 'Automated data download'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 11. WEB APPLICATION
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "11. Web Application", level=1)
    add_body(doc,
        "A complete Flask web application provides real-time Alzheimer's detection with user authentication, "
        "MRI image upload, model selection, confidence scoring, and detailed clinical context including "
        "probable causes and treatment suggestions."
    )

    add_heading_styled(doc, "Key Features", level=2)
    add_bullet(doc, "User Registration & Login with secure PBKDF2 password hashing.")
    add_bullet(doc, "MRI Image Upload supporting JPG/PNG formats (max 16MB).")
    add_bullet(doc, "Model Selection from 6 trained deep learning architectures.")
    add_bullet(doc, "Real-time Prediction with per-class probability distribution.")
    add_bullet(doc, "Clinical Context with detailed probable causes and treatment suggestions.")
    add_bullet(doc, "Prediction History stored in SQLite database.")
    add_bullet(doc, "Responsive Design with Bootstrap 5 for mobile compatibility.")

    add_heading_styled(doc, "Application Routes", level=2)
    add_styled_table(doc,
        ['Route', 'Description'],
        [
            ['/', 'Home page with project statistics'],
            ['/about', 'Disease information, dataset, technology stack'],
            ['/register', 'New user registration'],
            ['/login', 'User authentication'],
            ['/predict', 'MRI upload and prediction (requires login)'],
            ['/methodology', 'ML pipeline and training documentation'],
        ]
    )

    add_heading_styled(doc, "Database Schema", level=2)
    add_styled_table(doc,
        ['Table', 'Field', 'Type', 'Description'],
        [
            ['users', 'id', 'INTEGER PK', 'Auto-increment primary key'],
            ['users', 'full_name', 'TEXT', 'User full name'],
            ['users', 'email', 'TEXT UNIQUE', 'User email address'],
            ['users', 'password_hash', 'TEXT', 'PBKDF2-SHA256 hashed password'],
            ['predictions', 'id', 'INTEGER PK', 'Auto-increment primary key'],
            ['predictions', 'user_id', 'INTEGER FK', 'References users.id'],
            ['predictions', 'image_filename', 'TEXT', 'Uploaded MRI filename'],
            ['predictions', 'model_used', 'TEXT', 'Model architecture used'],
            ['predictions', 'predicted_class', 'TEXT', 'Predicted Alzheimer stage'],
            ['predictions', 'confidence', 'REAL', 'Prediction confidence (%)'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 12. CONCLUSION
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "12. Conclusion", level=1)

    add_heading_styled(doc, "Key Findings", level=2)
    add_bullet(doc, "Transfer learning with fine-tuning vastly outperforms training from scratch on medical imaging tasks.")
    add_bullet(doc, "MobileNetV2 with partial backbone unfreezing achieved 99.47% test accuracy with only 2.6M parameters.")
    add_bullet(doc, "Input resolution matters: upgrading from 128x128 to 224x224 improved accuracy by +24.20 percentage points.")
    add_bullet(doc, "Models trained from scratch (CNN, MLP) completely failed to converge on this imbalanced dataset.")
    add_bullet(doc, "The fine-tuned model achieved perfect 1.00 F1-score on the rarest class (Moderate Dementia, only 73 test samples).")
    add_bullet(doc, "Class weighting effectively addressed the 138:1 imbalance ratio between majority and minority classes.")

    add_heading_styled(doc, "Limitations", level=2)
    add_bullet(doc, "The model is trained on a single dataset (OASIS) and should be validated externally before clinical use.")
    add_bullet(doc, "MRI preprocessing (skull stripping, normalization) was not applied, which could further improve performance.")
    add_bullet(doc, "This system is for educational purposes only and is NOT a substitute for professional medical diagnosis.")

    add_heading_styled(doc, "Future Work", level=2)
    add_bullet(doc, "Implement Grad-CAM visualization to show which brain regions influence predictions.")
    add_bullet(doc, "Add support for 3D volumetric MRI analysis (currently using 2D slices).")
    add_bullet(doc, "Validate on external datasets (ADNI, MIRIAD) for clinical robustness.")
    add_bullet(doc, "Implement model ensembling and prediction confidence calibration.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 13. REFERENCES
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, "13. References", level=1)
    refs = [
        "[1] OASIS Dataset: https://www.oasis-brains.org/",
        "[2] Kaggle Dataset: https://www.kaggle.com/datasets/ninadaithal/imagesoasis",
        "[3] MobileNetV2: Sandler et al., 'MobileNetV2: Inverted Residuals and Linear Bottlenecks', CVPR 2018.",
        "[4] VGG16: Simonyan & Zisserman, 'Very Deep Convolutional Networks for Large-Scale Image Recognition', ICLR 2015.",
        "[5] ResNet50: He et al., 'Deep Residual Learning for Image Recognition', CVPR 2016.",
        "[6] EfficientNet: Tan & Le, 'EfficientNet: Rethinking Model Scaling for CNNs', ICML 2019.",
        "[7] TensorFlow: https://www.tensorflow.org/",
        "[8] Flask: https://flask.palletsprojects.com/",
        "[9] Alzheimer's Association: https://www.alz.org/",
        "[10] Mayo Clinic - Alzheimer's Treatments: https://www.mayoclinic.org/",
    ]
    for r in refs:
        add_body(doc, r)

    # ── Save ──
    output_path = os.path.join(SCRIPT_DIR, 'Alzheimer_Project_Report.docx')
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    print("Building editable Word document...")
    path = build_docx()
    print(f"\n{'='*60}")
    print(f"  Report generated: {path}")
    print(f"{'='*60}")
